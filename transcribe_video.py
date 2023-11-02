
#? Code to transcribe videos from any youtube URL


import argparse
import json
import os
import shutil
import subprocess
import time
from pathlib import Path

import validators
from docx import Document
from dotenv import find_dotenv, load_dotenv

from utils.database import DB_MANAGER
from utils.processing import transcribe_audio_whisper
from utils.utils import (clean_string, get_video_id_from_ytube_url,
                         send_discord_msg)
from utils.yt_utils import get_video_metadata, ic

#* load environment variables
load_dotenv(find_dotenv())

class Video:
    def __init__(self, input_args, config):
        self._config = config
        self.video_url = input_args.get("url")
        # times in seconds, iterations should not be greater than max time
        self.stats = {
            "run_time": 0,
            "start_time": time.time(),
            "iterations": 0,
            "transcriptions": [],
        }
        self.exit_on_video = input_args.get("exit_on_video", True)
        self.metadata = get_video_metadata(self.video_url)
        self.curr_errors = 0
        table_name = os.getenv("TABLE_NAME")
        if table_name is None:
            table_name = input_args.get("table_name")
        if table_name is None:
            table_name = self.get_channel_from_name().replace(" ", "_")
        # fallback to video_id
        if table_name == "" or table_name is None:
            try:
                ic("Grabbing table name from video_id")
                video_id = get_video_id_from_ytube_url(self.video_url)
                self.video_id = video_id
            except Exception as e:
                ic(e)
                ic("Error getting video id")
                self.video_id = ""
        else:
            ic(f"setting table_name from TABLE_NAME as: {table_name}")
            self.video_id = table_name
            
        self.video_title = clean_string(self.metadata.get("title", table_name))
        self.video_folder = os.path.join(Path(os.path.dirname(os.path.abspath(__file__))), "videos", self.video_title)
        
        try:
            folder_path = self.video_folder
            if os.path.exists(folder_path):
                shutil.rmtree(folder_path)
            os.makedirs(folder_path)
        except Exception as e:
            ic("Failed to make folder")


        try:
            self.db_manager = DB_MANAGER()
        except Exception as e:
            ic("Failed to make db_manager")
            print(e)
            self.db_manager = None
            # exit(1)

        try:
            print("Creating table for named: ", self.video_id)
            # create table if it doesn't exist
            self.db_manager.create_tables(self.video_id)
        except Exception as e:
            print(e)
            ic("Failed to make table")
            exit(1)

        self.save_to_db = input_args.get("save_to_db", True)

        self.transcribe_tool = input_args.get("transcribe_tool", "whisper")

    def get_channel_from_name(self):
        metadata = self.metadata
        channel = metadata.get("channel", "")
        uploader_id = metadata.get("uploader_id", "")
        uploader = metadata.get("uploader", "")
        if uploader == "Yahoo Finance":
            return "YahooFinance"
        if uploader_id == "Bloomberg":
            return "Bloomberg"
        return channel

    def download_file_from_url(self):
        print("downloading video", self.video_url)
        filename = os.path.join(self.video_folder, "video.mp4")
        result = subprocess.run(
            f'yt-dlp -f \"bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best\" -o {filename} --cookies ytcookies.txt {self.video_url}',
            shell=True, capture_output=True
        )
        ic(result)
        return result.stdout.decode("utf-8")


    def transcribe_video(self):
        # check if url is a local path or a url
        if validators.url(self.video_url):    
            filename = os.path.join(self.video_folder, "video.mp4")
            self.download_file_from_url()
        else:
            filename = self.video_url
        data = transcribe_audio_whisper(filename, False)
        partial_output = filename.replace(".mp4", ".json")

        with open(partial_output, "w", encoding="utf-8", errors="ignore") as f:
            f.write(json.dumps(data, indent=0))
        
        text = data.get("text", "")
            
        # write to doc file
        transcript_doc = Document()
        transcript_doc.add_heading(f'Video: {self.video_title}', 0)
        transcript_doc.add_paragraph(text)
        docx_output = filename.replace(".mp4", ".docx")
        transcript_doc.save(docx_output)

        ds_data = {
            "content": data.get("text", "")
        }
        # adjust runtime based on iteration if available
        # split data content into chunks of 1900 characters
        CHUNK_SIZE = 1900
        chunks = [ds_data['content'][i:i+CHUNK_SIZE] for i in range(0, len(ds_data['content']), CHUNK_SIZE)]
        for chunk in chunks:
            try:
                send_discord_msg({
                    "content": chunk,
                })
            except Exception as e:
                return

        
        # load file
        # load file
        pass

if __name__ == "__main__":
    # todo add argparser
    parser = argparse.ArgumentParser(description='Process video from url')
    parser.add_argument("--url", help="url where you can download video", default="https://www.dropbox.com/s/jvajeen28clpicy/archive.mp4?dl=1")
    parser.add_argument('--exit_for_videos', '-efv', help='exit for videos, or non livestreams', default=False)
    # save_to_db
    parser.add_argument('--save_to_db', '-stdb', help='save to db', default=True)
    parser.add_argument('--transcribe_tool', '-tt', help='select transcribe tool,\nwhisper = 0\nwit ai = 1', default=0)
    
    args = parser.parse_args()
    # ensure WIT_AI_TOKEN is set
    ic("Running main")
    if os.environ.get("WIT_AI_TOKEN") is None:
        print("WIT_AI_TOKEN is not set")
        # exit(1)
    dict_args = {
        "url": args.url,
        "exit_on_video": args.exit_for_videos,
        "save_to_db": args.save_to_db,
        "transcribe_tool": "whisper"
    }
    
    private_video = Video(dict_args, {})

    # todo add logging
    private_video.transcribe_video()
